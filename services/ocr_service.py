"""
OCR 服务模块 — PP-OCRv6 tiny (RapidOCR ONNX) 异步识别
"""
import os
import io
import time
import threading
import concurrent.futures
import re as _re
from pathlib import Path
from typing import Optional

from config import OUTPUT_DIR
from utils.logging_config import get_logger
from utils.cleanup import cleanup_file as _cleanup

logger = get_logger(__name__)

# ── OCR 任务存储 ──
# 任务状态：task_id -> {status, progress, total, message, error,
#                      src_path, dst_path, src_filename, started_at, cancel}
_ocr_tasks: dict = {}
_ocr_tasks_lock = threading.Lock()
_OCR_EXECUTOR = concurrent.futures.ThreadPoolExecutor(
    max_workers=1, thread_name_prefix='ocr'
)

# OCR 引擎单例
_ocr_engine = None
_ocr_engine_lock = threading.Lock()

# 输入限制
OCR_MAX_FILE_BYTES = 50 * 1024 * 1024  # 50MB
OCR_MAX_PAGES = 100


def _get_ocr_engine():
    """OCR 引擎单例。使用 RapidOCR (ONNX Runtime)"""
    global _ocr_engine
    if _ocr_engine is not None:
        return _ocr_engine
    with _ocr_engine_lock:
        if _ocr_engine is not None:
            return _ocr_engine
        from rapidocr_onnxruntime import RapidOCR
        _ocr_engine = RapidOCR()
        logger.info("OCR 引擎已初始化")
    return _ocr_engine


def _update_task(task_id: str, **kwargs):
    """线程安全地更新 OCR 任务状态"""
    with _ocr_tasks_lock:
        t = _ocr_tasks.get(task_id)
        if t:
            t.update(kwargs)


def get_task(task_id: str) -> Optional[dict]:
    """获取 OCR 任务状态"""
    with _ocr_tasks_lock:
        t = _ocr_tasks.get(task_id)
        if t:
            return dict(t)
        return None


def remove_task(task_id: str):
    """移除 OCR 任务"""
    with _ocr_tasks_lock:
        _ocr_tasks.pop(task_id, None)


def create_task(src_path: Path, dst_path: Path, src_filename: str) -> str:
    """创建 OCR 任务并返回 task_id"""
    import uuid as _uuid
    task_id = _uuid.uuid4().hex
    with _ocr_tasks_lock:
        _ocr_tasks[task_id] = {
            'status': 'pending',
            'progress': 0,
            'total': 0,
            'message': '已入队，等待识别...',
            'error': None,
            'src_path': src_path,
            'dst_path': dst_path,
            'src_filename': src_filename,
            'started_at': time.time(),
            'cancel': False,
        }
    _OCR_EXECUTOR.submit(_ocr_worker, task_id, src_path, dst_path)
    logger.info(f"OCR 任务已创建: {task_id}, 文件: {src_filename}")
    return task_id


def cleanup_orphaned_tasks():
    """清理过期 OCR 任务"""
    cutoff = time.time() - 1800  # 30 分钟
    with _ocr_tasks_lock:
        expired = [tid for tid, t in _ocr_tasks.items()
                   if t.get('started_at', 0) < cutoff]
        for tid in expired:
            t = _ocr_tasks.pop(tid, None)
            if t:
                _cleanup(t.get('src_path'))
                _cleanup(t.get('dst_path'))
                logger.debug(f"已清理过期 OCR 任务: {tid}")


# ── OCR 版式还原算法 ──────────────────────────────────

def _ocr_group_lines(boxes, texts, y_tol_ratio=0.5):
    """把 (box, text) 列表按 Y 中心聚成行"""
    if not boxes:
        return []
    items = list(zip(boxes, texts))
    items.sort(key=lambda x: ((x[0][1] + x[0][3]) / 2, x[0][0]))
    lines = []
    current = [items[0]]
    for it in items[1:]:
        y_center = (it[0][1] + it[0][3]) / 2
        avg_h = sum(b[3] - b[1] for b, _ in current) / len(current)
        last_y = (current[-1][0][1] + current[-1][0][3]) / 2
        if abs(y_center - last_y) < max(avg_h * y_tol_ratio, 5):
            current.append(it)
        else:
            current.sort(key=lambda x: x[0][0])
            lines.append(current)
            current = [it]
    if current:
        current.sort(key=lambda x: x[0][0])
        lines.append(current)
    return lines


def _ocr_split_paragraphs(lines, gap_ratio=0.7):
    """根据行间 Y 间距把 lines 切分成段落"""
    if not lines:
        return []
    paragraphs = [[lines[0]]]
    for prev, curr in zip(lines, lines[1:]):
        prev_bottom = max(b[3] for b, _ in prev)
        curr_top = min(b[1] for b, _ in curr)
        gap = curr_top - prev_bottom
        prev_h = sum(b[3] - b[1] for b, _ in prev) / len(prev)
        curr_h = sum(b[3] - b[1] for b, _ in curr) / len(curr)
        avg_h = (prev_h + curr_h) / 2
        if avg_h > 0 and gap > avg_h * gap_ratio:
            paragraphs.append([curr])
        else:
            paragraphs[-1].append(curr)
    return paragraphs


def _ocr_format_line(line):
    """行内多个文本块按 X 间距智能加空格"""
    if len(line) == 1:
        return line[0][1]
    parts = []
    for i, (box, text) in enumerate(line):
        if i > 0:
            prev_x2 = line[i - 1][0][2]
            curr_x1 = box[0]
            gap = curr_x1 - prev_x2
            w = box[2] - box[0]
            n = max(len(text), 1)
            char_w = max(w / n, 8)
            n_space = max(1, int(gap / char_w)) if gap >= char_w * 0.5 else 1
            parts.append(' ' * n_space)
        parts.append(text)
    return ''.join(parts)


def _ocr_avg_line_height(lines):
    """计算全页平均行高"""
    all_h = [b[3] - b[1] for line in lines for b, _ in line]
    return sum(all_h) / len(all_h) if all_h else 20


_HEADING_PATTERNS = [
    _re.compile(r'第[一二三四五六七八九十百\d]+[页章节]'),
    _re.compile(r'【.+】'),
    _re.compile(r'^([一二三四五六七八九十]+)、'),
    _re.compile(r'^\d+[\.、]\s*\S'),
]


def _ocr_classify_paragraph(para_lines, all_lines, para_idx, total_paras):
    """综合判断段落类型"""
    if not para_lines:
        return ('Normal', None)
    first_line = para_lines[0]
    first_text = _ocr_format_line(first_line).strip()
    avg_h = _ocr_avg_line_height(all_lines)
    line_h = sum(b[3] - b[1] for b, _ in first_line) / len(first_line)
    size_big = line_h > avg_h * 1.25
    all_y_min = min(b[1] for line in all_lines for b, _ in line)
    all_y_max = max(b[3] for line in all_lines for b, _ in line)
    para_y = first_line[0][0][1]
    in_top_quarter = (para_y - all_y_min) < (all_y_max - all_y_min) * 0.3
    is_h2_pattern = (
        any(p.search(first_text) for p in _HEADING_PATTERNS[:2])
        or bool(_HEADING_PATTERNS[2].match(first_text))
    )
    is_numbered = bool(_HEADING_PATTERNS[3].match(first_text))
    if size_big and in_top_quarter:
        return ('Heading 1', None)
    if is_h2_pattern:
        return ('Heading 2', None)
    if size_big:
        return ('Heading 1', None)
    if len(para_lines) == 1 and (
        first_text.startswith(('- ', '•', '·'))
        or is_numbered
        or _HEADING_PATTERNS[2].match(first_text)
    ):
        return ('List Bullet', None)
    if len(para_lines) > 1:
        list_count = sum(1 for ln in para_lines
                        if (ln[0][1].strip().startswith(('- ', '•', '·'))
                            or _HEADING_PATTERNS[3].match(ln[0][1].strip())
                            or _HEADING_PATTERNS[2].match(ln[0][1].strip())))
        if list_count >= max(2, len(para_lines) * 0.6):
            return ('List Bullet', None)
    return ('Normal', None)


def _ocr_is_table_row(line):
    """判断单行是否为表格行"""
    return len(line) >= 2


def _ocr_detect_tables_keep_order(paragraphs):
    """在段落列表中查找连续的表格行"""
    if not paragraphs:
        return [], [], []
    normal = []
    tables = []
    anchors = []
    current_table = []
    current_anchor = -1
    for para in paragraphs:
        is_table_row = (len(para) == 1 and _ocr_is_table_row(para[0]))
        if is_table_row:
            current_table.append(para)
        else:
            if current_table:
                col_counts = [len(p[0]) for p in current_table]
                if len(current_table) >= 2 and (max(col_counts) - min(col_counts) <= 1):
                    tables.append(current_table)
                    anchors.append(current_anchor)
                else:
                    normal.extend(current_table)
                current_table = []
            normal.append(para)
            current_anchor = len(normal) - 1
    if current_table:
        col_counts = [len(p[0]) for p in current_table]
        if len(current_table) >= 2 and (max(col_counts) - min(col_counts) <= 1):
            tables.append(current_table)
            anchors.append(current_anchor)
        else:
            normal.extend(current_table)
    return normal, tables, anchors


def _ocr_write_table(doc, table_paragraphs):
    """把表格段落写为 docx 表格"""
    from docx.shared import Pt
    n_rows = len(table_paragraphs)
    n_cols = max(len(p[0]) for p in table_paragraphs)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Light Grid Accent 1'
    for r_idx, para in enumerate(table_paragraphs):
        line = para[0]
        for c_idx, (box, text) in enumerate(line):
            if c_idx < n_cols:
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)


def _ocr_worker(task_id: str, src_path: Path, dst_path: Path):
    """OCR 后台任务（在独立线程中执行）"""
    try:
        _update_task(task_id, status='running', message='正在初始化 OCR 引擎...')
        ocr = _get_ocr_engine()
        from pdf2image import convert_from_path
        from docx import Document

        _update_task(task_id, message='正在解析 PDF 页码...')
        from pdf2image.pdf2image import pdfinfo_from_path
        try:
            info = pdfinfo_from_path(str(src_path))
            total_pages = info.get('Pages', 0)
        except Exception:
            total_pages = 0
        if total_pages > OCR_MAX_PAGES:
            _update_task(task_id, status='failed',
                         error=f'PDF 页数 {total_pages} 超过限制 {OCR_MAX_PAGES} 页')
            return
        if total_pages == 0:
            total_pages = 1

        _update_task(task_id, total=total_pages, progress=0,
                     message=f'正在识别 0/{total_pages} 页...')

        images = convert_from_path(str(src_path), dpi=200, fmt='jpeg',
                                   thread_count=2)

        doc = Document()
        from docx.shared import Pt
        style = doc.styles['Normal']
        style.font.name = 'DejaVu Sans'
        style.font.size = Pt(11)

        for i, img in enumerate(images):
            with _ocr_tasks_lock:
                t = _ocr_tasks.get(task_id)
                if t and t.get('cancel'):
                    _update_task(task_id, status='failed', error='用户已取消')
                    return

            tmp_img = dst_path.parent / f'_ocr_{task_id}_p{i}.jpg'
            img.convert('RGB').save(str(tmp_img), 'JPEG', quality=85)
            try:
                result = ocr(str(tmp_img))
            finally:
                _cleanup(tmp_img)

            page_items = []
            if result:
                detections = result[0] if len(result) >= 1 else None
                if detections:
                    for det in detections:
                        if not isinstance(det, (list, tuple)) or len(det) < 2:
                            continue
                        poly = det[0]
                        txt = det[1]
                        if not isinstance(poly, (list, tuple)) or len(poly) < 4:
                            continue
                        if not isinstance(txt, str) or not txt.strip():
                            continue
                        xs = [pt[0] for pt in poly]
                        ys = [pt[1] for pt in poly]
                        box = [min(xs), min(ys), max(xs), max(ys)]
                        page_items.append((box, txt.strip()))

            if i > 0:
                doc.add_page_break()

            if not page_items:
                doc.add_paragraph('')
            else:
                boxes = [it[0] for it in page_items]
                texts = [it[1] for it in page_items]
                lines = _ocr_group_lines(boxes, texts)
                paragraphs = _ocr_split_paragraphs(lines)
                normal_paras, table_groups, tbl_anchor = _ocr_detect_tables_keep_order(paragraphs)
                tbl_iter = iter(zip(table_groups, tbl_anchor))
                next_tbl = next(tbl_iter, None)
                for p_idx, para in enumerate(normal_paras):
                    while next_tbl and next_tbl[1] == p_idx:
                        _ocr_write_table(doc, next_tbl[0])
                        doc.add_paragraph('')
                        next_tbl = next(tbl_iter, None)
                    line_strs = [_ocr_format_line(line) for line in para]
                    text = '\n'.join(line_strs)
                    style, _ = _ocr_classify_paragraph(para, lines, p_idx, len(paragraphs))
                    if style == 'Normal':
                        doc.add_paragraph(text)
                    elif style == 'List Bullet':
                        try:
                            doc.add_paragraph(text, style='List Bullet')
                        except KeyError:
                            doc.add_paragraph(text)
                    else:
                        try:
                            doc.add_heading(text, level=1 if style == 'Heading 1' else 2)
                        except Exception:
                            doc.add_paragraph(text)
                while next_tbl:
                    _ocr_write_table(doc, next_tbl[0])
                    doc.add_paragraph('')
                    next_tbl = next(tbl_iter, None)

            _update_task(task_id, progress=i + 1,
                         message=f'正在识别 {i + 1}/{total_pages} 页...')

        doc.save(str(dst_path))
        _update_task(task_id, status='success',
                     message=f'识别完成，共 {total_pages} 页')
        logger.info(f"OCR 任务完成: {task_id}, {total_pages} 页")

    except ImportError as e:
        _update_task(task_id, status='failed',
                     error=f'缺少依赖库: {e}。请执行: pip install rapidocr_onnxruntime pdf2image Pillow')
        logger.error(f"OCR 任务 {task_id} 缺少依赖: {e}")
    except Exception as e:
        _update_task(task_id, status='failed', error=f'OCR 失败: {e}')
        logger.error(f"OCR 任务 {task_id} 失败: {e}")


# ── 初始化：启动时清理过期任务 ──
def init():
    """模块初始化"""
    cleanup_orphaned_tasks()
    logger.debug("OCR 服务模块已初始化")
